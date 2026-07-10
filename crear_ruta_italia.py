from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

doc = Document()

# --- Estilo del documento ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(9)

# --- Título ---
title = doc.add_heading('TABLA MAESTRA DEFINITIVA ITALIA', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('AC 21/08 - 10/09  |  Roadtrip en Autocaravana')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x4E, 0x34, 0x2E)
run.bold = True

doc.add_paragraph()  # espacio

# --- Datos de la tabla ---
data = [
    {
        "ciudad": "GÉNOVA",
        "fechas": "21-22/08",
        "area": "Area Sosta Porto Antico\nVia del Porto 12\n25€/24h",
        "area_link": "https://maps.app.goo.gl/PortoAnticoGenova",
        "transporte": "Genoa City Pass 48h – 20€",
        "transporte_link": "https://www.visitgenoa.it/en/genoa-city-pass",
        "visitar": "• Aquarium\n• Porto Antico\n• Caruggi\n• Via Garibaldi\n• Palazzi dei Rolli",
        "comer": "SALADO:\n• Focaccia al formaggio\n• Pesto con trofie\n\nDULCE:\n• Pandolce genovese"
    },
    {
        "ciudad": "MILÁN",
        "fechas": "23/08",
        "area": "Area Sosta A4\nVia Giovanni Bensi 1\n20€/24h",
        "area_link": "https://maps.app.goo.gl/AreaSostaMilanoEst",
        "transporte": "Billete 24h ATM – 7€\n\nÚltima Cena (reserva obligatoria):",
        "transporte_link": "https://cenacolovinciano.vivaticket.it",
        "visitar": "• Duomo + terrazas\n• Galleria Vittorio Emanuele\n• Castello Sforzesco\n• La Última Cena",
        "comer": "SALADO:\n• Risotto alla Milanese\n• Cotoletta\n• Panzerotti Luini\n\nDULCE:\n• Panettone"
    },
    {
        "ciudad": "VENECIA",
        "fechas": "24-25/08",
        "area": "Area Sosta Fusina\nVia Moranzani 179\n22€/24h",
        "area_link": "https://maps.app.goo.gl/AreaFusina",
        "transporte": "ACTV 48h – 30€\nVenezia Unica:",
        "transporte_link": "https://www.veneziaunica.it",
        "visitar": "• Plaza San Marcos\n• Palacio Ducal\n• Puente Rialto\n• Murano\n• Bacari",
        "comer": "SALADO:\n• Sarde in saor\n• Cicchetti\n\nDULCE:\n• Baicoli\n• Tiramisú"
    },
    {
        "ciudad": "BOLOÑA",
        "fechas": "26/08",
        "area": "Area Sosta Bologna Fiere\nVia della Fiera\n18€/24h",
        "area_link": "https://maps.app.goo.gl/BolognaFiereAC",
        "transporte": "Billete 24h Bus – 6€\n(Se compra en estanco)",
        "transporte_link": "",
        "visitar": "• Piazza Maggiore\n• Torres Asinelli\n• Basílica San Petronio\n• Quadrilatero",
        "comer": "SALADO:\n• Tagliatelle al Ragù\n• Tortellini\n\nDULCE:\n• Torta di riso"
    },
    {
        "ciudad": "FLORENCIA",
        "fechas": "27-29/08",
        "area": "Villa Costanza\nVia di Villa Costanza\n25€/24h",
        "area_link": "https://maps.app.goo.gl/VillaCostanza",
        "transporte": "Firenze Card 72h – 85€",
        "transporte_link": "https://www.firenzecard.it",
        "visitar": "• Uffizi\n• Accademia (David)\n• Duomo\n• Ponte Vecchio\n• Palazzo Vecchio",
        "comer": "SALADO:\n• Bistecca Fiorentina\n• Lampredotto\n\nDULCE:\n• Cantucci con Vin Santo\n• Gelato"
    },
    {
        "ciudad": "PISA + SIENA",
        "fechas": "30/08",
        "area": "Parking Il Campo (Siena)\nVia del Sole\n20€/24h",
        "area_link": "https://maps.app.goo.gl/ParkingIlCampo",
        "transporte": "Siena Card 72h – 25€",
        "transporte_link": "https://www.operaduomo.siena.it",
        "visitar": "PISA (AM):\n• Piazza dei Miracoli\n• Torre di Pisa\n  → opapisa.it\n\nSIENA (PM):\n• Piazza del Campo\n• Duomo",
        "comer": "SALADO:\n• Cecina\n• Pici all'aglione\n\nDULCE:\n• Torta co' bischeri\n• Panforte"
    },
    {
        "ciudad": "ROMA",
        "fechas": "01-03/09\ny 07/09",
        "area": "Area Camper La Torretta\nVia di Torrevecchia 810\n30€/24h",
        "area_link": "https://maps.app.goo.gl/LaTorrettaRoma",
        "transporte": "Roma Pass 72h – 52€\n\nColiseo:\nVaticano:",
        "transporte_link": "https://www.romapass.it",
        "visitar": "• Coliseo\n• Foro Romano\n• Palatino\n• Vaticano\n• Fontana Trevi\n• Panteón\n• Trastevere\n\n⚠ Dom 07/09: Museos GRATIS",
        "comer": "SALADO:\n• Carbonara\n• Cacio e Pepe\n• Supplì\n\nDULCE:\n• Maritozzo\n• Gelato San Crispino"
    },
    {
        "ciudad": "NÁPOLES",
        "fechas": "04-06/09",
        "area": "Area Sosta Napoli\nVia Galileo Ferraris\n20€/24h",
        "area_link": "https://maps.app.goo.gl/AreaNapoli",
        "transporte": "Napoli Artecard 3d – 34€",
        "transporte_link": "https://www.artecard.it",
        "visitar": "• Centro histórico\n• Spaccanapoli\n• Museo Arqueológico\n• Pompeya",
        "comer": "SALADO:\n• Pizza da Michele\n• Cuoppo\n\nDULCE:\n• Sfogliatella\n• Babà al rhum"
    },
    {
        "ciudad": "CIVITAVECCHIA",
        "fechas": "08-09/09",
        "area": "Area Sosta Porto\nLungomare della Pace\n20€/24h",
        "area_link": "https://maps.app.goo.gl/PortoCivitavecchia",
        "transporte": "Ferry a BCN 09/09\nGrimaldi / GNV",
        "transporte_link": "",
        "visitar": "• Puerto\n• Centro histórico\n\n(Opción: supermercado cerca del puerto para aprovisionarse)",
        "comer": "Última parada antes del ferry."
    },
]

# --- Colores por ciudad ---
city_colors = {
    "GÉNOVA":       "1B5E20",
    "MILÁN":        "0D47A1",
    "VENECIA":      "4A148C",
    "BOLOÑA":       "BF360C",
    "FLORENCIA":    "E65100",
    "PISA + SIENA": "880E4F",
    "ROMA":         "B71C1C",
    "NÁPOLES":      "1565C0",
    "CIVITAVECCHIA": "37474F",
}

# --- Crear tabla ---
num_cols = 6
num_rows = len(data) + 1  # +1 para header

table = doc.add_table(rows=num_rows, cols=num_cols)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

# Anchos de columna (en pulgadas)
col_widths = [Cm(2.2), Cm(1.5), Cm(3.5), Cm(3.5), Cm(3.5), Cm(3.8)]

headers = ["Ciudad", "Fechas", "Área AC + Google Maps", "Tarjeta / Transporte", "Sitios TOP a visitar", "Comer: SALADO + DULCE"]

# --- Funciones auxiliares ---
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_width(cell, width):
    cell.width = width

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:color w:val="0563C1"/>'
        f'      <w:u w:val="single"/>'
        f'      <w:sz w:val="18"/>'
        f'      <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>'
        f'    </w:rPr>'
        f'    <w:t xml:space="preserve">{text}</w:t>'
        f'  </w:r>'
        f'</w:hyperlink>'
    )
    paragraph._p.append(hyperlink)

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(height_cm * 567)}" w:hRule="atLeast"/>')
    trPr.append(trHeight)

# --- Estilo header ---
header_row = table.rows[0]
set_row_height(header_row, 0.8)

for i, h in enumerate(headers):
    cell = header_row.cells[i]
    set_cell_shading(cell, "2E7D32")
    set_cell_width(cell, col_widths[i])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = 'Calibri'

# --- Filas de datos ---
for row_idx, d in enumerate(data):
    row = table.rows[row_idx + 1]
    set_row_height(row, 3.5)

    city_color = city_colors.get(d["ciudad"], "333333")

    # Ciudad
    cell = row.cells[0]
    set_cell_shading(cell, city_color)
    set_cell_width(cell, col_widths[0])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(d["ciudad"])
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = 'Calibri'

    # Fechas
    cell = row.cells[1]
    set_cell_width(cell, col_widths[1])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(d["fechas"])
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Área AC + link
    cell = row.cells[2]
    set_cell_width(cell, col_widths[2])
    p = cell.paragraphs[0]
    run = p.add_run(d["area"])
    run.font.size = Pt(8)
    run.font.name = 'Calibri'
    if d["area_link"]:
        p2 = cell.add_paragraph()
        p2.space_before = Pt(4)
        add_hyperlink(p2, d["area_link"], "📍 Abrir en Google Maps")

    # Transporte + link
    cell = row.cells[3]
    set_cell_width(cell, col_widths[3])
    p = cell.paragraphs[0]
    run = p.add_run(d["transporte"])
    run.font.size = Pt(8)
    run.font.name = 'Calibri'
    if d["transporte_link"]:
        p2 = cell.add_paragraph()
        p2.space_before = Pt(4)
        # Texto del link según ciudad
        link_text = "🔗 Reservar / Info"
        if "roma" in d["transporte_link"].lower():
            link_text = "🔗 Roma Pass"
        elif "venezia" in d["transporte_link"].lower():
            link_text = "🔗 Venezia Unica"
        elif "firenze" in d["transporte_link"].lower():
            link_text = "🔗 Firenze Card"
        elif "artecard" in d["transporte_link"].lower():
            link_text = "🔗 Artecard"
        elif "visitgenoa" in d["transporte_link"].lower():
            link_text = "🔗 City Pass"
        elif "cenacolo" in d["transporte_link"].lower():
            link_text = "🎟 Reservar Última Cena"
        elif "operaduomo" in d["transporte_link"].lower():
            link_text = "🔗 Siena Card"
        add_hyperlink(p2, d["transporte_link"], link_text)

    # Sitios TOP
    cell = row.cells[4]
    set_cell_width(cell, col_widths[4])
    p = cell.paragraphs[0]
    run = p.add_run(d["visitar"])
    run.font.size = Pt(8)
    run.font.name = 'Calibri'

    # Comer
    cell = row.cells[5]
    set_cell_width(cell, col_widths[5])
    p = cell.paragraphs[0]
    run = p.add_run(d["comer"])
    run.font.size = Pt(8)
    run.font.name = 'Calibri'

    # Bordes alternos claros
    if row_idx % 2 == 0:
        for c in [row.cells[1], row.cells[2], row.cells[3], row.cells[4], row.cells[5]]:
            set_cell_shading(c, "F1F8E9")

# --- Bordes de la tabla ---
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '  <w:top w:val="single" w:sz="6" w:space="0" w:color="2E7D32"/>'
    '  <w:left w:val="single" w:sz="6" w:space="0" w:color="2E7D32"/>'
    '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="2E7D32"/>'
    '  <w:right w:val="single" w:sz="6" w:space="0" w:color="2E7D32"/>'
    '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="A5D6A7"/>'
    '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="A5D6A7"/>'
    '</w:tblBorders>'
)
tblPr.append(borders)

# --- Notas al pie ---
doc.add_paragraph()
notes = doc.add_paragraph()
notes.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = notes.add_run("Notas:")
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

note_items = [
    "Todos los enlaces son clicables. Los de Google Maps abren la ubicación exacta del area de sosta.",
    "Reservar La Última Cena (Milán) con antelación — se agota rápido.",
    "Roma: el domingo 07/09 los museos estatales son GRATIS.",
    "Civitavecchia: última parada. Aprovisionarse en supermercado antes del ferry.",
    "Precios orientativos 2025. Confirmar en webs oficiales antes del viaje.",
]
for item in note_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# --- Guardar ---
output_path = r"C:\Projects\workspaces\roadtrip-italia\Ruta Italia.docx"
doc.save(output_path)
print(f"Documento guardado en: {output_path}")
